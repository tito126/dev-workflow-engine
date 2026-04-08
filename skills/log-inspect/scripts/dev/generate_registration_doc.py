#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成参赛报名 Word 文档"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

def create_registration_doc(output_path):
    doc = Document()
    
    # 设置默认字体
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    doc.styles['Normal'].font.size = Pt(11)
    
    # 标题
    title = doc.add_heading('技术氛围挑战赛报名材料', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 副标题
    subtitle = doc.add_paragraph('AI 驱动的生产日志智能巡检')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(14)
    
    doc.add_paragraph()
    
    # 一、团队信息
    doc.add_heading('一、团队信息', level=1)
    
    # 团队名称
    p = doc.add_paragraph()
    p.add_run('团队名称：').bold = True
    p.add_run('[请填写]')
    
    # 成员表格
    table = doc.add_table(rows=3, cols=4)
    table.style = 'Table Grid'
    
    # 表头
    headers = ['角色', '姓名', '部门', '职位']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    # 队长行
    table.rows[1].cells[0].text = '队长'
    table.rows[1].cells[1].text = '[请填写]'
    table.rows[1].cells[2].text = '[请填写]'
    table.rows[1].cells[3].text = '[请填写]'
    
    # 成员行
    table.rows[2].cells[0].text = '成员'
    table.rows[2].cells[1].text = '[请填写]'
    table.rows[2].cells[2].text = '[请填写]'
    table.rows[2].cells[3].text = '[请填写]'
    
    doc.add_paragraph()
    
    # 二、参赛方向
    doc.add_heading('二、参赛方向', level=1)
    p = doc.add_paragraph()
    p.add_run('问题排查 — AI 驱动的生产日志智能巡检').bold = True
    
    doc.add_paragraph()
    
    # 三、初步方案
    doc.add_heading('三、初步方案', level=1)
    
    # 1. 痛点分析
    doc.add_heading('1. 痛点分析', level=2)
    
    pain_points = [
        ('日志获取繁琐', '需手动开端口、登录服务器、定位文件、下载，流程长、易出错'),
        ('分析效率低', '单个日志文件动辄 200MB+，人工翻看耗时数小时，依赖个人经验'),
        ('问题发现滞后', '通常等用户反馈才介入排查，缺乏主动巡检机制'),
        ('多环境割裂', '普通服务器（gz 文件）和 K8s（Loki）分析方式不同，无法统一'),
    ]
    
    for title, desc in pain_points:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{title}：').bold = True
        p.add_run(desc)
    
    doc.add_paragraph()
    
    # 2. AI 解决方案
    doc.add_heading('2. AI 解决方案', level=2)
    
    p = doc.add_paragraph()
    p.add_run('通过 AI 实现"企微对话触发 → 自动拉取日志 → 智能分析 → 生成报告"的全流程自动化：')
    
    # 解决方案表格
    table2 = doc.add_table(rows=5, cols=3)
    table2.style = 'Table Grid'
    
    headers2 = ['环节', 'AI 能力', '工具/方法']
    for i, header in enumerate(headers2):
        table2.rows[0].cells[i].text = header
        table2.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    solutions = [
        ('交互触发', '企业微信对话触发，自然语言下达巡检指令', '企微机器人 + AI Agent'),
        ('日志预处理', '自动提取 ERROR/WARN、识别慢接口、错误分类聚合', 'Python 脚本'),
        ('智能分析', '理解日志语义、关联 traceId、定位根因、给出修复建议', '大模型（Claude/GPT）'),
        ('报告生成', '结构化输出：统计概览 + 详细分析 + 慢接口 + 改进建议', 'AI 生成 HTML 报告'),
    ]
    
    for i, (col1, col2, col3) in enumerate(solutions):
        table2.rows[i+1].cells[0].text = col1
        table2.rows[i+1].cells[1].text = col2
        table2.rows[i+1].cells[2].text = col3
    
    doc.add_paragraph()
    
    # 3. 预期价值
    doc.add_heading('3. 预期价值', level=2)
    
    # 价值表格
    table3 = doc.add_table(rows=5, cols=4)
    table3.style = 'Table Grid'
    
    headers3 = ['指标', '传统方式', 'AI 巡检', '提效']
    for i, header in enumerate(headers3):
        table3.rows[0].cells[i].text = header
        table3.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    values = [
        ('日志获取', '30 分钟+', '1 分钟', '30x'),
        ('200MB 日志分析', '2~4 小时', '5 分钟', '30x'),
        ('问题发现', '被动响应', '主动巡检', '风险前置'),
        ('跨环境支持', '两套流程', '统一抽象', '降低门槛'),
    ]
    
    for i, (col1, col2, col3, col4) in enumerate(values):
        table3.rows[i+1].cells[0].text = col1
        table3.rows[i+1].cells[1].text = col2
        table3.rows[i+1].cells[2].text = col3
        table3.rows[i+1].cells[3].text = col4
    
    doc.add_paragraph()
    
    # 已验证案例
    p = doc.add_paragraph()
    p.add_run('已验证案例：').bold = True
    p.add_run('浙江省中（普通服务器 200MB 日志）、乐山市人民医院（K8s Loki 日志）均已跑通，成功生成巡检报告并发现多项生产问题。')
    
    # 保存
    doc.save(output_path)
    print(f'文档已生成: {output_path}')

if __name__ == '__main__':
    create_registration_doc('参赛报名_AI日志智能巡检.docx')
